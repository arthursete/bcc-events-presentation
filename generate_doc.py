from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1.27)
    # Light grey background via shading
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F2F2F2'
    })
    shading.append(shd)
    return p


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Digital Solutions IA1 — Technical Proposal', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('BCC Events Finder', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub2 = add_para('A Web Application for Brisbane City Council Event Venues')
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info = add_para('Kedron State High School, 2026')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()


# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading('1. Introduction', level=1)

add_heading('Introduction to the Presentation', level=2)
add_para('This technical proposal presents the design and planning of BCC Events Finder, a purpose-built web application that transforms the Brisbane City Council Events Locations dataset into an accessible, user-friendly venue discovery platform. The application addresses a clear gap in the current open data ecosystem: while Brisbane City Council publishes comprehensive venue data through their Open Data Portal, this information is presented in a raw, technical format that is inaccessible to everyday residents and visitors.')
add_para('This proposal will be presented as a multimodal slideshow with annotated wireframes, data dictionary tables, pseudocode algorithms, and SQL query demonstrations. The presentation is targeted at Brisbane City Council stakeholders — including data managers, IT decision-makers, and community engagement officers — using appropriate technical language.')

add_heading('Constraints', level=2)

add_para('Time Constraints', bold=True)
add_bullet('The project proposal must be completed within approximately 15 hours of class time')
add_bullet('Final submission due 06/03/2026')
add_bullet('Draft submission required by Term 1 Week 4')

add_para('Data Constraints', bold=True)
add_bullet('The dataset is sourced from Brisbane City Council\'s Open Data Portal in CSV/Excel format')
add_bullet('The dataset contains 1,696 venue location records across 186 suburbs')
add_bullet('Data includes 8 fields: Venue Name, Venue Address, Venue Type, Latitude, Longitude, Suburb, suburb_list, and geolocation')
add_bullet('Some venues have multiple venue types stored as comma-separated values, requiring normalisation')
add_bullet('The dataset is static (venue locations) and does not include live event scheduling data')

add_para('Legal Constraints', bold=True)
add_bullet('The dataset is published under a Creative Commons Attribution 4.0 (CC BY 4.0) licence, permitting free reuse with attribution')
add_bullet('Admin user credentials must be stored securely using password hashing (never plain text)')
add_bullet('The application must comply with the Australian Privacy Act 1988 regarding any personal data collected through admin accounts')
add_bullet('Brisbane City Council must be attributed as the data source')

add_heading('Prescribed and Self-Determined Criteria', level=2)

add_para('Prescribed Criteria (from Task Sheet)', bold=True)
add_bullet('The solution must be a web application')
add_bullet('It must allow users to access and filter the BCC Events Locations dataset')
add_bullet('It must present data in an easy-to-understand format')
add_bullet('It must allow effective searching of the dataset')
add_bullet('It must include administrative privileges where relevant')
add_bullet('The proposal must be presented as a multimodal presentation for a technical audience')

add_para('Self-Determined Criteria', bold=True)

add_para('Personal Impact:', bold=True)
add_bullet('The application should empower individual residents to discover venues in their own suburb, encouraging personal engagement with community spaces and reducing the barrier to participation in local events')

add_para('Social Impact:', bold=True)
add_bullet('By making venue data accessible to all users regardless of technical ability, the application promotes social inclusion and community connectedness. Accessibility features (clear labels, high contrast, logical navigation) ensure users with disabilities can access the information')

add_para('Economic Impact:', bold=True)
add_bullet('A user-friendly venue discovery platform can boost local economic activity by driving foot traffic to community venues, markets, and cultural centres. Tourism visitors can more easily find attractions, benefiting local businesses in those areas')

doc.add_page_break()

# ============================================================
# 2. USER STORY
# ============================================================
add_heading('2. User Story', level=1)

add_heading('2.1 Description of Web Application', level=2)
add_para('BCC Events Finder is a proposed web application that provides Brisbane residents and visitors with an intuitive way to discover, search, and filter the 1,696 event venue locations maintained by Brisbane City Council. The application transforms raw open data into a visually engaging, easy-to-navigate experience.')

add_para('Key Features for Users (Public):', bold=True)
add_bullet('Home page with a prominent search bar and quick filter buttons for venue types')
add_bullet('Browse/search page with sidebar filters for suburb (186 options via searchable/autocomplete dropdown) and venue type (10 categories), sort options, and paginated results')
add_bullet('Venue detail page showing full venue information including address, type, suburb, and a map placeholder displaying GPS coordinates')
add_bullet('Related venues section showing other venues in the same suburb')

add_para('Key Features for Admins:', bold=True)
add_bullet('Secure login with username and password authentication (hashed passwords)')
add_bullet('Dashboard showing statistics (total venues, suburbs, venue types) and a searchable venue management table')
add_bullet('Add/Edit venue form with fields for venue name, address, suburb (dropdown), venue type (multi-select checkboxes), latitude, and longitude')
add_bullet('Delete venue functionality with confirmation prompt')

add_para('Proposed Software Stack:', bold=True)
add_table(
    ['Component', 'Technology', 'Justification'],
    [
        ['Frontend', 'HTML, CSS, JavaScript', 'Standard web technologies; accessible in all browsers'],
        ['Backend', 'Python (Flask)', 'Lightweight framework suitable for data-driven apps; widely taught'],
        ['Database', 'SQLite', 'File-based relational database; no server setup required; ideal for the dataset size'],
        ['Mapping', 'Leaflet.js (future)', 'Open-source mapping library compatible with the lat/long data'],
    ]
)

doc.add_paragraph()

add_heading('2.2 Existing Solution', level=2)
add_para('Brisbane City Council Open Data Portal (data.brisbane.qld.gov.au)', bold=True)
add_para('The existing solution is the Brisbane City Council Open Data Portal, powered by the Opendatasoft platform. This is the council\'s official data publishing website where they host 384+ datasets, including the Events Locations dataset.')

add_para('The portal provides multiple views of the data:')
add_bullet('Table view: A paginated spreadsheet showing all 1,696 records with raw field values')
add_bullet('Map view: An interactive Leaflet.js map with venue markers and geographic shape-drawing tools for spatial filtering')
add_bullet('Analyse view: A basic chart builder for creating simple visualisations')
add_bullet('Export/API: Download options in CSV, JSON, Excel, and GeoJSON formats')

add_para('UI Components Used:', bold=True)
add_table(
    ['Component', 'Description'],
    [
        ['Top Navigation Bar', 'Links to Home, Stories, Data, Visualise, API, About'],
        ['Search Bar', 'Text-based keyword search across datasets and within datasets'],
        ['Theme Filters', 'Category filters on the catalogue page (Arts, Parks, Transport, etc.)'],
        ['Dataset Cards', 'Title, publisher, and update timestamp for each dataset'],
        ['Data Table', 'Paginated, sortable spreadsheet-style table'],
        ['Interactive Map', 'Leaflet.js map with markers and shape-drawing spatial filters'],
        ['Filter Sidebar', 'Faceted filter panel with active filter display'],
        ['Download Buttons', 'Export in multiple formats'],
    ]
)

doc.add_paragraph()
add_para('Pros of Existing Solution:', bold=True)
add_bullet('Comprehensive raw data access — Full dataset available for download in multiple formats')
add_bullet('API access — Programmatic access for developers via the Opendatasoft Explore API')
add_bullet('Geographic visualisation — Built-in map view with coordinate plotting')
add_bullet('Open licence — CC BY 4.0 allows free reuse')
add_bullet('Multiple export formats — CSV, JSON, Excel, GeoJSON')

add_para('Cons of Existing Solution:', bold=True)
add_bullet('Not designed for general public — The portal is a data catalogue for developers and researchers, not a venue discovery tool for everyday users')
add_bullet('No suburb-based browsing — Users cannot simply type a suburb name and see all nearby venues in a friendly layout')
add_bullet('No user-friendly venue type filtering — Filters show raw field values rather than visual categories with icons')
add_bullet('No contextual information — Venue records are displayed as raw data rows without additional context such as nearby amenities, accessibility details, or venue descriptions')
add_bullet('Raw data presentation — Shows latitude/longitude, database field names, and unformatted text with no contextual information')
add_bullet('No personalisation — No saved favourites, recommendations, or accessibility preferences')
add_bullet('Poor mobile experience — The desktop-focused data exploration interface does not work well on mobile devices')
add_bullet('No proximity search — Users cannot search by distance from their current location')
add_bullet('Overwhelming for non-technical users — API documentation, JSON schemas, and metadata clutter the interface')

doc.add_page_break()

# ============================================================
# 3. DATA STORY
# ============================================================
add_heading('3. Data Story', level=1)

add_heading('3.1 Sample Data', level=2)

add_para('Dataset Overview:', bold=True)
add_table(
    ['Property', 'Detail'],
    [
        ['Source', 'Brisbane City Council Open Data Portal (data.brisbane.qld.gov.au)'],
        ['Dataset Name', 'Brisbane City Council Events Locations'],
        ['File Format', 'CSV / Excel (.xlsx)'],
        ['Licence', 'Creative Commons Attribution 4.0 (CC BY 4.0)'],
        ['Total Records', '1,696 venue locations'],
        ['Total Columns', '8 fields'],
        ['Unique Suburbs', '186'],
        ['Unique Venue Types', '10'],
        ['Accessibility', 'Publicly available; no authentication required to download'],
        ['Security', 'Read-only public dataset; admin features in the proposed app will use hashed password authentication'],
    ]
)

doc.add_paragraph()
add_para('Dataset Columns:', bold=True)
add_table(
    ['Column', 'Data Type', 'Example'],
    [
        ['Venue Name', 'Text', '"7th Brigade Park, Chermside"'],
        ['Venue address', 'Text', '"7th Brigade Park, 21 Kittyhawk Drive, Chermside"'],
        ['Venue type', 'Text (may be comma-separated)', '"Park" or "Community Garden,Park"'],
        ['Latitude', 'Decimal', '-27.3794'],
        ['Longitude', 'Decimal', '153.0303'],
        ['Suburb', 'Text', '"Chermside"'],
        ['suburb_list', 'Text', '"C, Chermside"'],
        ['geolocation', 'Text', '"-27.3794, 153.0303"'],
    ]
)

doc.add_paragraph()
add_para('Sample Data (First 5 Records):', bold=True)
add_table(
    ['Venue Name', 'Venue Address', 'Venue Type', 'Lat', 'Long', 'Suburb'],
    [
        ["'Grow for Life' Community Garden", '41 Abbeville St, Upper Mt Gravatt', 'Community Garden, Non-Council venue', '-27.544', '153.089', 'Upper Mt Gravatt'],
        ['7th Brigade Park, Chermside', '21 Kittyhawk Drive, Chermside', 'Park', '-27.379', '153.030', 'Chermside'],
        ['A.J. Jones Recreation Reserve', '247 Cornwall St, Greenslopes', 'Park', '-27.503', '153.044', 'Greenslopes'],
        ['A.R.C. Hill Park, Virginia', '93 Goss Road, Virginia', 'Park', '-27.384', '153.056', 'Virginia'],
        ['Abbeville Street Park', '41 Abbeville St, Upper Mount Gravatt', 'Community Garden, Park', '-27.545', '153.090', 'Upper Mt Gravatt'],
    ]
)

doc.add_paragraph()
add_para('Venue Type Distribution:', bold=True)
add_table(
    ['Venue Type', 'Approximate Count'],
    [
        ['Park', '~1,200+ (majority)'],
        ['Library', '~33'],
        ['Community Hall', '~50'],
        ['Swimming Pool', '~24'],
        ['Community Garden', '~40'],
        ['General Council venue', '~30'],
        ['Non-Council venue', '~20'],
        ['Golf course', '~10'],
        ['Light up asset', '~15'],
        ['Mobile library', '~5'],
    ]
)
add_para('Note: Some venues have multiple types (e.g., "Community Garden, Park"), requiring a many-to-many relationship in the normalised database.', italic=True)

doc.add_paragraph()

add_heading('3.2 Data Dictionary', level=2)
add_para('The original flat dataset is normalised into five relational tables to eliminate redundancy and support efficient querying.')

add_para('Table 1: Suburbs', bold=True)
add_para('Stores the 186 unique Brisbane suburbs.')
add_table(
    ['Field', 'Data Type', 'Constraints', 'Description'],
    [
        ['SuburbID', 'INT', 'PRIMARY KEY, AUTO_INCREMENT', 'Unique identifier for each suburb'],
        ['SuburbName', 'VARCHAR(50)', 'NOT NULL, UNIQUE', 'Name of the suburb (e.g. "Chermside")'],
    ]
)

doc.add_paragraph()
add_para('Table 2: VenueTypes', bold=True)
add_para('Stores the 10 venue type categories.')
add_table(
    ['Field', 'Data Type', 'Constraints', 'Description'],
    [
        ['TypeID', 'INT', 'PRIMARY KEY, AUTO_INCREMENT', 'Unique identifier for each type'],
        ['TypeName', 'VARCHAR(50)', 'NOT NULL, UNIQUE', 'Name of the venue type (e.g. "Park")'],
    ]
)

doc.add_paragraph()
add_para('Table 3: Venues', bold=True)
add_para('Stores the 1,696 venue records with location data.')
add_table(
    ['Field', 'Data Type', 'Constraints', 'Description'],
    [
        ['VenueID', 'INT', 'PRIMARY KEY, AUTO_INCREMENT', 'Unique identifier for each venue'],
        ['VenueName', 'VARCHAR(100)', 'NOT NULL', 'Name of the venue'],
        ['VenueAddress', 'VARCHAR(150)', 'NOT NULL', 'Full street address'],
        ['Latitude', 'DECIMAL(10,7)', 'NOT NULL', 'GPS latitude (-27.65 to -27.28)'],
        ['Longitude', 'DECIMAL(10,7)', 'NOT NULL', 'GPS longitude (152.79 to 153.42)'],
        ['SuburbID', 'INT', 'FK REFERENCES Suburbs(SuburbID)', 'Links venue to its suburb'],
    ]
)

doc.add_paragraph()
add_para('Table 4: VenueTypeLink (Junction Table)', bold=True)
add_para('Resolves the many-to-many relationship between venues and types.')
add_table(
    ['Field', 'Data Type', 'Constraints', 'Description'],
    [
        ['VenueID', 'INT', 'FK REFERENCES Venues(VenueID), COMPOSITE PK', 'Links to the venue'],
        ['TypeID', 'INT', 'FK REFERENCES VenueTypes(TypeID), COMPOSITE PK', 'Links to the type'],
    ]
)

doc.add_paragraph()
add_para('Table 5: Users', bold=True)
add_para('Stores admin accounts for content management.')
add_table(
    ['Field', 'Data Type', 'Constraints', 'Description'],
    [
        ['UserID', 'INT', 'PRIMARY KEY, AUTO_INCREMENT', 'Unique identifier'],
        ['Username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', 'Admin login username'],
        ['PasswordHash', 'VARCHAR(255)', 'NOT NULL', 'Hashed password (never plain text)'],
        ['Role', 'VARCHAR(20)', "NOT NULL, DEFAULT 'admin'", 'Access control role'],
        ['CreatedDate', 'DATETIME', 'NOT NULL, DEFAULT CURRENT_TIMESTAMP', 'Account creation date'],
    ]
)

doc.add_paragraph()
add_para('Entity-Relationship Diagram:', bold=True)
add_code_block(
"""Suburbs (1) ────< (M) Venues (M) >──── (M) VenueTypes
                            \\                /
                             \\── VenueTypeLink ──/

Users (standalone — admin authentication)"""
)

add_bullet('Suburbs to Venues: One-to-many (one suburb contains many venues)')
add_bullet('Venues to VenueTypes: Many-to-many via VenueTypeLink junction table')
add_bullet('Users: Independent table for admin login')

doc.add_paragraph()
add_para('Business Rules and Validation Constraints', bold=True)
add_para('The following rules govern data integrity and input validation across the application.')
add_table(
    ['Rule', 'Description'],
    [
        ['Latitude range', 'Must be between -27.65 and -27.28 (Brisbane metro area)'],
        ['Longitude range', 'Must be between 152.79 and 153.42 (Brisbane metro area)'],
        ['Venue name', 'Cannot be blank; maximum 100 characters'],
        ['Venue address', 'Cannot be blank; maximum 150 characters'],
        ['Suburb reference', 'Must reference an existing SuburbID in the Suburbs table'],
        ['Venue type', 'Each venue must have at least one entry in VenueTypeLink'],
        ['Username', 'Must be unique; cannot be blank'],
        ['Password', 'Stored as a salted hash; never stored in plain text'],
    ]
)

doc.add_page_break()

# ============================================================
# 4. PROPOSED SOLUTION
# ============================================================
add_heading('4. Proposed Solution', level=1)

add_heading('4.1 HCI (Human-Computer Interaction)', level=2)

add_para('Wireframes', bold=True)
add_para('The following annotated wireframes illustrate the proposed user interface for each page of the application.')

# Insert wireframe screenshots with captions
wireframes = [
    ('wireframe_home.png',      'Wireframe 1: Home Page'),
    ('wireframe_browse.png',    'Wireframe 2: Browse / Search Results'),
    ('wireframe_detail.png',    'Wireframe 3: Venue Detail'),
    ('wireframe_login.png',     'Wireframe 4: Admin Login'),
    ('wireframe_dashboard.png', 'Wireframe 5: Admin Dashboard'),
    ('wireframe_form.png',      'Wireframe 6: Add / Edit Venue Form'),
]

screenshots_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'wireframe_screenshots')

for img_file, caption in wireframes:
    img_path = os.path.join(screenshots_dir, img_file)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        # Centre the image paragraph
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.bold = True
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = 'Calibri'
        doc.add_paragraph()  # spacing between wireframes

doc.add_paragraph()
add_para('Design Principles Applied:', bold=True)
add_table(
    ['Principle', 'Application'],
    [
        ['Contrast', 'Dark navigation bar against light content creates visual separation; blue primary buttons contrast with white backgrounds; red delete buttons differentiate destructive actions from safe ones'],
        ['Alignment', 'Left-aligned form labels above inputs for fast vertical scanning; consistent table column alignment; two-column form layout for related fields'],
        ['Repetition', 'Consistent venue card design (image, name, suburb, type badge, link) across all listing pages; identical stat card structure on admin dashboard; same navigation bar on every page'],
        ['Proximity', 'Related filter controls grouped together in the sidebar; venue details and map placed side by side; suburb and venue type form fields adjacent; statistics cards clustered as a summary row'],
        ['Visual Hierarchy', 'Large hero heading draws attention first; venue name uses largest heading on detail page; stat values use large bold text on dashboard; breadcrumbs in smaller text for secondary navigation'],
    ]
)

doc.add_paragraph()
add_para('Usability Principles Applied:', bold=True)
add_table(
    ['Principle', 'Feature'],
    [
        ['Consistent navigation', 'Same navigation bar on every public page; distinct admin navigation on all admin pages'],
        ['Prominent search', 'Search bar placed in the hero section of the home page — the primary user task is immediately accessible'],
        ['User feedback', 'Results count ("Showing 1,696 of 1,696 venues") keeps users informed of filter results'],
        ['Breadcrumb navigation', 'Shows users their location in the site hierarchy and enables easy backtracking'],
        ['Pagination', 'Breaks 1,696 results into manageable pages to prevent information overload'],
        ['Related content', 'Venue detail page shows other venues in the same suburb for further discovery'],
        ['Placeholder guidance', 'Form inputs include example values (e.g., "-27.3794") to guide correct data entry'],
        ['Accessibility', 'Required fields marked with asterisks and legend; clear labels on all inputs for screen readers; high contrast text throughout'],
        ['Destructive action differentiation', 'Red delete button visually warns of permanent action vs. blue edit button'],
    ]
)

doc.add_paragraph()
add_para('User-Centred Design Considerations', bold=True)
add_para('The following table links specific interface design decisions to the needs of different user types who will interact with the application.')
add_table(
    ['User Type', 'Design Decision', 'Rationale'],
    [
        ['Older residents', 'High contrast text, large click targets, clear labels', 'Reduces barriers for users with reduced vision or limited digital literacy'],
        ['Mobile users', 'Responsive card layout, collapsible sidebar filters', 'Ensures usability on smaller screens without losing functionality'],
        ['Non-technical users', 'Category filter buttons with descriptive labels instead of raw field names', 'Makes venue types instantly understandable without database knowledge'],
        ['Brisbane visitors', 'Suburb-based browsing with searchable dropdown', 'Allows discovery by location even without prior knowledge of Brisbane geography'],
        ['Admin users', 'Distinct navy navigation bar, confirmation prompts on delete', 'Clearly separates admin context and prevents accidental destructive actions'],
    ]
)

doc.add_paragraph()
add_para('Note: The suburb dropdown (186 suburbs) will be implemented as a searchable/autocomplete dropdown rather than a plain select list. This allows users to type a suburb name and see matching suggestions, significantly improving usability over scrolling through a long list.', italic=True)

doc.add_paragraph()

add_heading('4.2 Algorithms', level=2)

add_para('SQL Queries', bold=True)

add_para('Create Tables:', bold=True)
add_code_block("""CREATE TABLE Suburbs (
    SuburbID   INTEGER PRIMARY KEY AUTOINCREMENT,
    SuburbName VARCHAR(50) NOT NULL UNIQUE
);

CREATE TABLE VenueTypes (
    TypeID   INTEGER PRIMARY KEY AUTOINCREMENT,
    TypeName VARCHAR(50) NOT NULL UNIQUE
);

CREATE TABLE Venues (
    VenueID      INTEGER PRIMARY KEY AUTOINCREMENT,
    VenueName    VARCHAR(100) NOT NULL,
    VenueAddress VARCHAR(150) NOT NULL,
    Latitude     DECIMAL(10,7) NOT NULL,
    Longitude    DECIMAL(10,7) NOT NULL,
    SuburbID     INT NOT NULL,
    FOREIGN KEY (SuburbID) REFERENCES Suburbs(SuburbID)
);

CREATE TABLE VenueTypeLink (
    VenueID INT NOT NULL,
    TypeID  INT NOT NULL,
    PRIMARY KEY (VenueID, TypeID),
    FOREIGN KEY (VenueID) REFERENCES Venues(VenueID),
    FOREIGN KEY (TypeID)  REFERENCES VenueTypes(TypeID)
);

CREATE TABLE Users (
    UserID       INTEGER PRIMARY KEY AUTOINCREMENT,
    Username     VARCHAR(50) NOT NULL UNIQUE,
    PasswordHash VARCHAR(255) NOT NULL,
    Role         VARCHAR(20) NOT NULL DEFAULT 'admin',
    CreatedDate  TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
);""")

add_para('Note: SQLite uses AUTOINCREMENT (one word) rather than MySQL\'s AUTO_INCREMENT. The TEXT type is used broadly in SQLite; VARCHAR constraints are shown here to document intended field lengths.', italic=True)

doc.add_paragraph()
add_para('Select venues in a specific suburb (JOIN):', bold=True)
add_code_block("""SELECT v.VenueID, v.VenueName, v.VenueAddress, v.Latitude, v.Longitude,
       s.SuburbName
FROM Venues v
INNER JOIN Suburbs s ON v.SuburbID = s.SuburbID
WHERE s.SuburbName = 'Chermside';""")

doc.add_paragraph()
add_para('Select venues filtered by type (JOIN through junction table):', bold=True)
add_code_block("""SELECT v.VenueID, v.VenueName, v.VenueAddress, s.SuburbName
FROM Venues v
INNER JOIN VenueTypeLink vtl ON v.VenueID = vtl.VenueID
INNER JOIN VenueTypes vt ON vtl.TypeID = vt.TypeID
INNER JOIN Suburbs s ON v.SuburbID = s.SuburbID
WHERE vt.TypeName = 'Library';""")

doc.add_paragraph()
add_para('Search venues by name (LIKE):', bold=True)
add_code_block("""SELECT v.VenueID, v.VenueName, v.VenueAddress, s.SuburbName
FROM Venues v
INNER JOIN Suburbs s ON v.SuburbID = s.SuburbID
WHERE v.VenueName LIKE '%Brigade%';""")

doc.add_paragraph()
add_para('Count venues per suburb (GROUP BY):', bold=True)
add_code_block("""SELECT s.SuburbName, COUNT(v.VenueID) AS VenueCount
FROM Suburbs s
LEFT JOIN Venues v ON s.SuburbID = v.SuburbID
GROUP BY s.SuburbName
ORDER BY VenueCount DESC;""")

doc.add_paragraph()
add_para('Insert a new venue (Admin):', bold=True)
add_code_block("""INSERT INTO Venues (VenueName, VenueAddress, Latitude, Longitude, SuburbID)
VALUES ('New Community Centre', '10 Example Street, Paddington', -27.4598, 153.0137, 42);

INSERT INTO VenueTypeLink (VenueID, TypeID)
VALUES (last_insert_rowid(), 3);""")

doc.add_paragraph()
add_para('Update venue details:', bold=True)
add_code_block("""UPDATE Venues
SET VenueName = 'Updated Venue Name',
    VenueAddress = '99 New Address Street, Ashgrove'
WHERE VenueID = 15;""")

doc.add_paragraph()
add_para('Delete a venue (with referential integrity):', bold=True)
add_code_block("""DELETE FROM VenueTypeLink WHERE VenueID = 15;
DELETE FROM Venues WHERE VenueID = 15;""")

doc.add_paragraph()
add_para('Pseudocode', bold=True)

add_para('Search Algorithm:', bold=True)
add_code_block("""FUNCTION SearchVenues(searchTerm)
    SANITISE searchTerm to prevent SQL injection

    IF searchTerm is empty THEN
        RETURN all venues (default listing)
    END IF

    QUERY database (parameterised):
        SELECT DISTINCT v.VenueID, v.VenueName, v.VenueAddress,
               s.SuburbName, vt.TypeName
        FROM Venues v
        JOIN Suburbs s ON v.SuburbID = s.SuburbID
        JOIN VenueTypeLink vtl ON v.VenueID = vtl.VenueID
        JOIN VenueTypes vt ON vtl.TypeID = vt.TypeID
        WHERE v.VenueName LIKE ?
           OR s.SuburbName LIKE ?
           OR vt.TypeName LIKE ?
        BIND "%" + searchTerm + "%" as parameter for each ?

    SET results TO query results

    RETURN results
END FUNCTION""")

doc.add_paragraph()
add_para('Filter Algorithm (Multi-Criteria):', bold=True)
add_code_block("""FUNCTION FilterVenues(selectedSuburb, selectedType, pageNumber, pageSize)
    SET baseQuery TO "SELECT DISTINCT v.VenueID, v.VenueName, v.VenueAddress,
                      s.SuburbName, vt.TypeName
                      FROM Venues v
                      JOIN Suburbs s ON v.SuburbID = s.SuburbID
                      JOIN VenueTypeLink vtl ON v.VenueID = vtl.VenueID
                      JOIN VenueTypes vt ON vtl.TypeID = vt.TypeID"
    SET conditions TO empty list
    SET parameters TO empty list

    IF selectedSuburb is NOT empty THEN
        ADD "s.SuburbName = ?" to conditions
        BIND selectedSuburb as parameter (not string concatenation)
    END IF

    IF selectedType is NOT empty THEN
        ADD "vt.TypeName = ?" to conditions
        BIND selectedType as parameter (not string concatenation)
    END IF

    IF conditions is NOT empty THEN
        APPEND " WHERE " + JOIN conditions WITH " AND " TO baseQuery
    END IF

    APPEND " ORDER BY v.VenueName ASC" TO baseQuery
    APPEND " LIMIT pageSize OFFSET (pageNumber - 1) * pageSize" TO baseQuery

    EXECUTE baseQuery with bound parameters
    SET results TO query results

    IF results is empty THEN
        DISPLAY "No venues match your filters"
    END IF

    RETURN results
END FUNCTION""")

doc.add_paragraph()
add_para('Sort Algorithm:', bold=True)
add_code_block("""FUNCTION SortVenues(venueList, sortField)
    IF sortField = "name" THEN
        SORT venueList BY VenueName in ascending order
    ELSE IF sortField = "suburb" THEN
        SORT venueList BY SuburbName in ascending order,
            THEN BY VenueName in ascending order
    END IF

    RETURN venueList
END FUNCTION""")

doc.add_paragraph()
add_para('Authentication Algorithm:', bold=True)
add_code_block("""FUNCTION AuthenticateUser(inputUsername, inputPassword)
    SANITISE inputUsername to prevent SQL injection

    QUERY database (parameterised):
        SELECT UserID, Username, PasswordHash, Role
        FROM Users
        WHERE Username = ?
        BIND inputUsername as parameter

    IF no record found THEN
        DISPLAY "Invalid username or password"
        RETURN false
    END IF

    SET storedHash TO query result PasswordHash

    IF VerifyPassword(inputPassword, storedHash) = true THEN
        CREATE session for user with UserID and Role
        SET session expiry to 30 minutes
        DISPLAY "Login successful"
        RETURN true
    ELSE
        INCREMENT failed login counter for inputUsername
        DISPLAY "Invalid username or password"
        RETURN false
    END IF
END FUNCTION""")

add_para('Note: VerifyPassword uses a salted hash comparison (e.g., bcrypt or werkzeug.security.check_password_hash) rather than manually re-hashing, because salted hashing produces different outputs each time. Additional security considerations include session expiration (auto-logout after inactivity) and failed login attempt handling (rate limiting or temporary lockout after repeated failures).', italic=True)

doc.add_paragraph()
add_para('Add Venue (CRUD - Create):', bold=True)
add_code_block("""FUNCTION AddVenue(name, address, latitude, longitude, suburbName, typeNames)
    CHECK user is logged in as admin
    IF NOT admin THEN
        DISPLAY "Access denied"
        RETURN
    END IF

    VALIDATE that name, address, latitude, longitude are not empty
    IF validation fails THEN
        DISPLAY "All fields are required"
        RETURN
    END IF

    VALIDATE latitude is between -27.65 and -27.28 (Brisbane metro range)
    VALIDATE longitude is between 152.79 and 153.42 (Brisbane metro range)
    IF coordinate validation fails THEN
        DISPLAY "Coordinates must be within the Brisbane metro area"
        RETURN
    END IF

    VALIDATE at least one typeName is selected
    IF typeNames is empty THEN
        DISPLAY "At least one venue type must be selected"
        RETURN
    END IF

    BEGIN TRANSACTION

    QUERY database to find SuburbID for suburbName
    IF suburb not found THEN
        INSERT new suburb into Suburbs table
        SET suburbID TO new suburb's ID
    ELSE
        SET suburbID TO existing suburb's SuburbID
    END IF

    INSERT INTO Venues (VenueName, VenueAddress, Latitude, Longitude, SuburbID)
    VALUES (name, address, latitude, longitude, suburbID)

    SET newVenueID TO last inserted ID

    FOR EACH typeName IN typeNames
        QUERY database to find TypeID for typeName
        INSERT INTO VenueTypeLink (VenueID, TypeID) VALUES (newVenueID, TypeID)
    END FOR

    COMMIT TRANSACTION (rollback on error)

    DISPLAY "Venue added successfully"
END FUNCTION""")

doc.add_paragraph()
add_para('Delete Venue (CRUD - Delete):', bold=True)
add_code_block("""FUNCTION DeleteVenue(venueID)
    CHECK user is logged in as admin
    IF NOT admin THEN
        DISPLAY "Access denied"
        RETURN
    END IF

    DISPLAY "Are you sure you want to delete this venue?"

    IF user confirms THEN
        BEGIN TRANSACTION
        DELETE FROM VenueTypeLink WHERE VenueID = venueID
        DELETE FROM Venues WHERE VenueID = venueID
        COMMIT TRANSACTION (rollback on error)
        DISPLAY "Venue deleted successfully"
    ELSE
        DISPLAY "Deletion cancelled"
    END IF
END FUNCTION""")

add_para('Note: Both DELETE statements are wrapped in a transaction to maintain referential integrity — if either statement fails, neither change is applied.', italic=True)

doc.add_paragraph()

add_heading('4.3 Navigation', level=2)
add_para('Site Map / User Flow:', bold=True)
add_code_block("""PUBLIC USER FLOW:
                    ┌──────────────┐
                    │   HOME PAGE  │
                    │  (Search +   │
                    │   Filters)   │
                    └──────┬───────┘
                           │
              ┌────────────┼────────────┐
              v            v            v
    ┌─────────────┐ ┌───────────┐ ┌──────────┐
    │   BROWSE /  │ │   ABOUT   │ │  ADMIN   │
    │   SEARCH    │ │   PAGE    │ │  LOGIN   │
    │  RESULTS    │ └───────────┘ └────┬─────┘
    └──────┬──────┘                    │
           │                           v
           v                    ┌──────────┐
    ┌─────────────┐             │  ADMIN   │
    │   VENUE     │             │DASHBOARD │
    │   DETAIL    │             └────┬─────┘
    │  (+ Related │                  │
    │   Venues)   │         ┌────────┼────────┐
    └─────────────┘         v        v        v
                        ┌──────┐ ┌──────┐ ┌──────┐
                        │ ADD  │ │ EDIT │ │DELETE│
                        │VENUE │ │VENUE │ │VENUE│
                        └──────┘ └──────┘ └──────┘""")

doc.add_paragraph()
add_para('Navigation Description:', bold=True)
add_bullet('Home Page is the entry point. Users can search, use quick filters, or browse featured venues')
add_bullet('From Home, users navigate to Browse/Search Results via the search bar or filter buttons')
add_bullet('From Browse, users click a venue card to view the Venue Detail page')
add_bullet('The Venue Detail page shows Related Venues in the same suburb, allowing further exploration without returning to search')
add_bullet('The Admin Login page is accessible from the navigation bar')
add_bullet('After successful login, admins are directed to the Admin Dashboard')
add_bullet('From the Dashboard, admins can Add, Edit, or Delete venues')
add_bullet('All pages include a consistent navigation bar for direct access to any section')

doc.add_paragraph()

add_heading('4.4 Acknowledgements', level=2)
add_bullet('Brisbane City Council. (2026). Events Locations Dataset. Brisbane City Council Open Data Portal. Retrieved February 2026, from https://data.brisbane.qld.gov.au/explore/dataset/brisbane-city-council-events-locations/')
add_bullet('Brisbane City Council. (2026). Open Data Portal. https://data.brisbane.qld.gov.au/')
add_bullet('Creative Commons. (n.d.). Attribution 4.0 International (CC BY 4.0). https://creativecommons.org/licenses/by/4.0/')
add_bullet('Opendatasoft. (n.d.). Opendatasoft Platform. https://www.opendatasoft.com/')

# Save
output_path = '/Users/arthursete/Digi tech 26/IA1_Technical_Proposal.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
